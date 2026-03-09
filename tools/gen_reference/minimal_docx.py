#!/usr/bin/env python3
"""
Generate minimal DOCX file for reference.

This script creates a simple "Hello World" DOCX file using python-docx.
The generated file is used to understand the XML structure and relationships.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def generate_test_images():
    """Generate simple test PNG images for reference docs (optional)."""
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        print("Pillow not installed — skipping test image generation")
        print("Install with: pip install Pillow")
        return False

    assets_dir = os.path.join(os.path.dirname(__file__),
                              "../../tests/fixtures/reference_docs/assets")
    os.makedirs(assets_dir, exist_ok=True)

    # Inline test image (200x150, blue with text)
    img = Image.new("RGB", (200, 150), color=(70, 130, 180))
    draw = ImageDraw.Draw(img)
    draw.text((50, 65), "inline.png", fill="white")
    img.save(os.path.join(assets_dir, "inline.png"))

    # Floating test image (300x200, green with text)
    img2 = Image.new("RGB", (300, 200), color=(60, 150, 60))
    draw2 = ImageDraw.Draw(img2)
    draw2.text((80, 90), "floating.png", fill="white")
    img2.save(os.path.join(assets_dir, "floating.png"))

    print(f"Created test images in: {assets_dir}")
    return True


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph (renders as page number in Word)."""
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    run._r.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_hyperlink(paragraph, url, text):
    """Add a real w:hyperlink element to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url,
                          "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_minimal_docx():
    """Create a minimal DOCX with basic content."""
    doc = Document()

    # Add a simple paragraph
    p = doc.add_paragraph("Hello World")

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__),
                               "../../tests/fixtures/reference_docs/minimal.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_layout_docx():
    """Create a DOCX with headers/footers, page and section breaks, and block quotes."""
    doc = Document()

    def add_toc_paragraph(document):
        """Insert a TOC field for levels 1-3 (auto-updatable in Word)."""
        p = document.add_paragraph()
        r = p.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\"1-3\\" \\h \\z \\u '
        r._r.append(instr)

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r._r.append(fld_sep)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r._r.append(fld_end)

    # Header and footer
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header: Reference Document"
    section.footer.paragraphs[0].text = "Footer: Page "
    add_page_number_field(section.footer.paragraphs[0])

    # Page break
    doc.add_paragraph("Page 1 content before page break.")
    doc.add_page_break()
    doc.add_paragraph("Page 2 content after page break.")

    # Section break (next page)
    doc.add_section(start_type=WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Section 2 begins here.")

    # Block quote
    block_quote = doc.add_paragraph("This is a block quote sample.")
    try:
        block_quote.style = doc.styles["Intense Quote"]
    except KeyError:
        block_quote.paragraph_format.left_indent = Inches(0.5)
        block_quote.paragraph_format.right_indent = Inches(0.5)
        for run in block_quote.runs:
            run.italic = True

    # Table of contents field (levels 1-3)
    doc.add_paragraph("Table of Contents (update field in Word):")
    add_toc_paragraph(doc)

    output_path = os.path.join(os.path.dirname(__file__), "../../tests/fixtures/reference_docs/layout.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_images_docx():
    """Create a DOCX with inline and floating images."""
    doc = Document()

    # Inline image
    doc.add_paragraph("Inline image below:")
    inline_img_path = os.path.join(os.path.dirname(__file__), "../../tests/fixtures/reference_docs/assets/inline.png")
    os.makedirs(os.path.dirname(inline_img_path), exist_ok=True)
    # Placeholder: caller should place an image at inline.png; keep reference for structure
    if os.path.exists(inline_img_path):
        doc.add_picture(inline_img_path, width=Inches(2))
    else:
        doc.add_paragraph("[inline.png missing — place test asset here]")

    # Floating image (left, wrapped)
    doc.add_paragraph("Floating image (left, tight wrap):")
    float_img_path = os.path.join(os.path.dirname(__file__), "../../tests/fixtures/reference_docs/assets/floating.png")
    if os.path.exists(float_img_path):
        pic = doc.add_picture(float_img_path, width=Inches(2.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # python-docx lacks native floating support; note in text for XML inspection
        doc.add_paragraph("(Floating positioning to be adjusted in XML if needed)")
    else:
        doc.add_paragraph("[floating.png missing — place test asset here]")

    output_path = os.path.join(os.path.dirname(__file__), "../../tests/fixtures/reference_docs/images.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_tables_docx():
    """Create a DOCX with various table configurations including merged cells."""
    doc = Document()

    doc.add_heading("Simple Table", level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for i, text in enumerate(["Col A", "Col B", "Col C"]):
        table.rows[0].cells[i].text = text
    for r in range(1, 3):
        for c in range(3):
            table.rows[r].cells[c].text = f"Row {r} Col {c}"

    doc.add_heading("Horizontal Merge", level=2)
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].merge(table2.rows[0].cells[2])
    table2.rows[0].cells[0].text = "Merged Header (spans 3 columns)"
    for r in range(1, 3):
        for c in range(3):
            table2.rows[r].cells[c].text = f"R{r}C{c}"

    doc.add_heading("Vertical Merge", level=2)
    table3 = doc.add_table(rows=4, cols=3)
    table3.style = "Table Grid"
    for c in range(3):
        table3.rows[0].cells[c].text = f"Header {c}"
    table3.rows[1].cells[0].merge(table3.rows[3].cells[0])
    table3.rows[1].cells[0].text = "Spans 3 rows"
    for r in range(1, 4):
        table3.rows[r].cells[1].text = f"R{r}C1"
        table3.rows[r].cells[2].text = f"R{r}C2"

    doc.add_heading("Styled Table", level=2)
    table4 = doc.add_table(rows=3, cols=3)
    table4.style = "Light Grid Accent 1"
    for c, h in enumerate(["Name", "Value", "Status"]):
        table4.rows[0].cells[c].text = h
    table4.rows[1].cells[0].text = "Alpha"
    table4.rows[1].cells[1].text = "100"
    table4.rows[1].cells[2].text = "Active"
    table4.rows[2].cells[0].text = "Beta"
    table4.rows[2].cells[1].text = "200"
    table4.rows[2].cells[2].text = "Inactive"

    output_path = os.path.join(os.path.dirname(__file__),
                               "../../tests/fixtures/reference_docs/tables.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_lists_docx():
    """Create a DOCX with various list configurations including nested lists."""
    doc = Document()

    doc.add_heading("Bullet List", level=2)
    doc.add_paragraph("Apple", style="List Bullet")
    doc.add_paragraph("Banana", style="List Bullet")
    doc.add_paragraph("Cherry", style="List Bullet")

    doc.add_heading("Numbered List", level=2)
    doc.add_paragraph("First", style="List Number")
    doc.add_paragraph("Second", style="List Number")
    doc.add_paragraph("Third", style="List Number")

    doc.add_heading("Nested Bullet List", level=2)
    doc.add_paragraph("Parent item 1", style="List Bullet")
    doc.add_paragraph("Child item 1a", style="List Bullet 2")
    doc.add_paragraph("Child item 1b", style="List Bullet 2")
    doc.add_paragraph("Parent item 2", style="List Bullet")
    doc.add_paragraph("Child item 2a", style="List Bullet 2")

    doc.add_heading("Nested Numbered List", level=2)
    doc.add_paragraph("First", style="List Number")
    doc.add_paragraph("First-A", style="List Number 2")
    doc.add_paragraph("First-B", style="List Number 2")
    doc.add_paragraph("Second", style="List Number")
    doc.add_paragraph("Second-A", style="List Number 2")

    doc.add_heading("Mixed Nested List", level=2)
    doc.add_paragraph("Parent One", style="List Number")
    doc.add_paragraph("Child bullet", style="List Bullet 2")
    doc.add_paragraph("Child bullet two", style="List Bullet 2")

    output_path = os.path.join(os.path.dirname(__file__),
                               "../../tests/fixtures/reference_docs/lists.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_code_blocks_docx():
    """Create a DOCX with code block styling."""
    doc = Document()

    doc.add_heading("Code Block Example", level=2)
    doc.add_paragraph("Here is some inline code and a code block:")

    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.5)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(6)
    run = code.add_run('fn main() {\n    println!("Hello, Lontar!");\n}')
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    code2 = doc.add_paragraph()
    code2.paragraph_format.left_indent = Inches(0.5)
    run2 = code2.add_run("SELECT * FROM users WHERE active = true;")
    run2.font.name = "Courier New"
    run2.font.size = Pt(10)

    output_path = os.path.join(os.path.dirname(__file__),
                               "../../tests/fixtures/reference_docs/code_blocks.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def create_styled_docx():
    """Create a DOCX with various text styles."""
    doc = Document()

    # Heading
    doc.add_heading("Document Title", level=1)

    # Paragraph with mixed formatting
    p = doc.add_paragraph()
    p.add_run("Bold text").bold = True
    p.add_run(" and ")
    p.add_run("italic text").italic = True
    p.add_run(" and ")
    p.add_run("underlined text").underline = True

    # Colored text
    p = doc.add_paragraph()
    run = p.add_run("Red colored text")
    run.font.color.rgb = RGBColor(255, 0, 0)

    # Different font sizes
    p = doc.add_paragraph()
    run = p.add_run("Small text (8pt)")
    run.font.size = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run("Large text (16pt)")
    run.font.size = Pt(16)

    # Paragraph alignment
    p = doc.add_paragraph("Left aligned (default)")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph("Center aligned")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Right aligned")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Bullet list
    doc.add_paragraph("First item", style='List Bullet')
    doc.add_paragraph("Second item", style='List Bullet')
    doc.add_paragraph("Third item", style='List Bullet')

    # Numbered list
    doc.add_paragraph("First", style='List Number')
    doc.add_paragraph("Second", style='List Number')
    doc.add_paragraph("Third", style='List Number')

    # Table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header 1'
    hdr_cells[1].text = 'Header 2'
    hdr_cells[2].text = 'Header 3'

    # Data rows
    table.rows[1].cells[0].text = 'Data 1'
    table.rows[1].cells[1].text = 'Data 2'
    table.rows[1].cells[2].text = 'Data 3'

    table.rows[2].cells[0].text = 'Data 4'
    table.rows[2].cells[1].text = 'Data 5'
    table.rows[2].cells[2].text = 'Data 6'

    # Strikethrough
    p = doc.add_paragraph()
    run = p.add_run("Strikethrough text")
    run.font.strike = True

    # Superscript
    p = doc.add_paragraph()
    p.add_run("E = mc")
    run = p.add_run("2")
    run.font.superscript = True

    # Subscript
    p = doc.add_paragraph()
    p.add_run("H")
    run = p.add_run("2")
    run.font.subscript = True
    p.add_run("O")

    # Different font families
    p = doc.add_paragraph()
    run = p.add_run("Arial font")
    run.font.name = "Arial"

    p = doc.add_paragraph()
    run = p.add_run("Times New Roman font")
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run("Courier New font")
    run.font.name = "Courier New"

    # Hyperlink
    p = doc.add_paragraph("Visit: ")
    add_hyperlink(p, "https://example.com", "Example Website")

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__),
                               "../../tests/fixtures/reference_docs/styled.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

def create_headings_docx():
    """Create a DOCX with all heading levels."""
    doc = Document()

    doc.add_heading("Heading Level 1", level=1)
    doc.add_paragraph("Body text under heading 1")

    doc.add_heading("Heading Level 2", level=2)
    doc.add_paragraph("Body text under heading 2")

    doc.add_heading("Heading Level 3", level=3)
    doc.add_paragraph("Body text under heading 3")

    doc.add_heading("Heading Level 4", level=4)
    doc.add_paragraph("Body text under heading 4")

    doc.add_heading("Heading Level 5", level=5)
    doc.add_paragraph("Body text under heading 5")

    doc.add_heading("Heading Level 6", level=6)
    doc.add_paragraph("Body text under heading 6")

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__),
                               "../../tests/fixtures/reference_docs/headings.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

if __name__ == "__main__":
    print("Generating DOCX reference documents...")
    generate_test_images()
    create_minimal_docx()
    create_styled_docx()
    create_headings_docx()
    create_tables_docx()
    create_lists_docx()
    create_code_blocks_docx()
    create_images_docx()
    create_layout_docx()
    print("Done! Generated 8 reference DOCX files.")
