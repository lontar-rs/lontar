#!/usr/bin/env python3
"""Generate missing DOCX reference documents."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def set_cell_background(cell, fill):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def generate_headers_footers():
    """Generate headers_footers.docx"""
    doc = Document()
    
    # Add header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Header text appears at the top of each page."
    
    # Add footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Footer text appears at the bottom of each page."
    
    # Add content
    doc.add_paragraph("Document with header and footer.")
    doc.add_paragraph("This page has a header and footer.")
    
    doc.save('tests/fixtures/reference_docs/headers_footers.docx')
    print("✓ headers_footers.docx")

def generate_page_breaks():
    """Generate page_breaks.docx"""
    doc = Document()
    
    doc.add_paragraph("Page 1 content.")
    add_page_break(doc)
    doc.add_paragraph("Page 2 content.")
    add_page_break(doc)
    doc.add_paragraph("Page 3 content.")
    
    doc.save('tests/fixtures/reference_docs/page_breaks.docx')
    print("✓ page_breaks.docx")

def generate_table_of_contents():
    """Generate table_of_contents.docx"""
    doc = Document()
    
    # Add TOC field (simplified - just add a heading structure)
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('Section 1', style='List Bullet')
    doc.add_paragraph('Section 2', style='List Bullet')
    doc.add_paragraph('Subsection 2.1', style='List Bullet 2')
    
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Content for section 1.')
    
    doc.add_heading('Section 2', level=1)
    doc.add_paragraph('Content for section 2.')
    
    doc.add_heading('Subsection 2.1', level=2)
    doc.add_paragraph('Content for subsection 2.1.')
    
    doc.save('tests/fixtures/reference_docs/table_of_contents.docx')
    print("✓ table_of_contents.docx")

def generate_hyperlinks():
    """Generate hyperlinks.docx"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    doc = Document()
    
    # Add hyperlink paragraph
    p = doc.add_paragraph('This is a ')
    
    # Create relationship for external hyperlink
    part = doc.part
    relId = part.relate_to('https://example.com', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create hyperlink XML with proper structure
    hyperlink_xml = (
        f'<w:hyperlink {nsdecls("w", "r")} r:id="{relId}">'
        '<w:r>'
        '<w:rPr>'
        '<w:rStyle w:val="Hyperlink"/>'
        '<w:u w:val="single"/>'
        '<w:color w:val="0563C1"/>'
        '</w:rPr>'
        '<w:t>hyperlink</w:t>'
        '</w:r>'
        '</w:hyperlink>'
    )
    hyperlink_element = parse_xml(hyperlink_xml)
    p._element.append(hyperlink_element)
    p.add_run(' to an external website.')
    
    # Add internal reference hyperlink (using anchor attribute instead of relationship)
    p2 = doc.add_paragraph('This is a ')
    
    hyperlink_xml2 = (
        f'<w:hyperlink {nsdecls("w")} w:anchor="section-1">'
        '<w:r>'
        '<w:rPr>'
        '<w:rStyle w:val="Hyperlink"/>'
        '<w:u w:val="single"/>'
        '<w:color w:val="0563C1"/>'
        '</w:rPr>'
        '<w:t>link to section</w:t>'
        '</w:r>'
        '</w:hyperlink>'
    )
    hyperlink_element2 = parse_xml(hyperlink_xml2)
    p2._element.append(hyperlink_element2)
    p2.add_run(' within the document.')
    
    doc.save('tests/fixtures/reference_docs/hyperlinks.docx')
    print("✓ hyperlinks.docx")

def generate_block_quotes():
    """Generate block_quotes.docx"""
    doc = Document()
    
    # Add block quote with indentation
    p1 = doc.add_paragraph('This is a block quote.', style='Quote')
    p1.paragraph_format.left_indent = Inches(0.5)
    
    p2 = doc.add_paragraph('It can span multiple lines.', style='Quote')
    p2.paragraph_format.left_indent = Inches(0.5)
    
    # Add nested quote
    p3 = doc.add_paragraph('This is a nested quote.', style='Quote')
    p3.paragraph_format.left_indent = Inches(0.5)
    
    p4 = doc.add_paragraph('Inner quote level 1', style='Quote')
    p4.paragraph_format.left_indent = Inches(1.0)
    
    p5 = doc.add_paragraph('Inner quote level 2', style='Quote')
    p5.paragraph_format.left_indent = Inches(1.5)
    
    doc.save('tests/fixtures/reference_docs/block_quotes.docx')
    print("✓ block_quotes.docx")

if __name__ == '__main__':
    generate_headers_footers()
    generate_page_breaks()
    generate_table_of_contents()
    generate_hyperlinks()
    generate_block_quotes()
    print("\nAll missing DOCX files generated!")
